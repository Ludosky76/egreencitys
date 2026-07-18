"""
EGREENCITY'S — Genere les 8 emails de LANCEMENT OFFICIEL au format .docx
========================================================================

Document Word entierement editable avec mise en forme charte EGREENCITY'S :
- Logo en en-tete
- Bandeau "LANCEMENT OFFICIEL" vert vif
- Sections avec titres verts
- Liste a puces avec check verts
- Tableau calendrier vert clair
- Encadres "Pourquoi commune ?" jaune
- Etapes numerotees avec ronds verts
- Bloc CTA vert fonce
- Signature + footer

Lit _data_mairies.json et produit 8 .docx personnalises par commune.
"""
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DATA_FILE = ROOT / "_dossiers" / "02_Dossier_ADVENIR" / "Courriers_Mairies" / "_data_mairies.json"
OUT_DIR = ROOT / "_dossiers" / "02_Dossier_ADVENIR" / "Emails_Mairies"
LOGO_PATH = ROOT / "logo.png"

# ============================================================
#  Charte EGREENCITY'S
# ============================================================
COLOR_GREEN_DARK   = RGBColor(0x0A, 0x48, 0x00)
COLOR_GREEN_MED    = RGBColor(0x2E, 0x7D, 0x32)
COLOR_GREEN_BRIGHT = RGBColor(0x33, 0xCC, 0x00)
COLOR_BLUE         = RGBColor(0x2B, 0x4D, 0xB5)
COLOR_TEXT         = RGBColor(0x33, 0x33, 0x33)
COLOR_TEXT_DARK    = RGBColor(0x1A, 0x3A, 0x00)
COLOR_GREY         = RGBColor(0x60, 0x7D, 0x8B)
COLOR_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_YELLOW_BG    = "FFF9C4"
COLOR_YELLOW_DARK  = RGBColor(0x5D, 0x40, 0x37)
COLOR_GREEN_BG     = "F4FFF0"
COLOR_GREEN_LIGHT  = "E8F5E9"
COLOR_HEADER_BG    = "0A4800"

FONT = "Calibri"


# ============================================================
#  Helpers
# ============================================================
def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def borders(cell, color="BFBFBF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def paragraph_shade(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def paragraph_border_left(p, color, sz=24):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def run_style(run, *, bold=False, italic=False, color=COLOR_TEXT, size=10.5, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def add_para(doc, text, *, bold=False, italic=False, color=COLOR_TEXT,
             size=10.5, before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run_style(run, bold=bold, italic=italic, color=color, size=size)
    return p


def add_runs_mixed(p, segments):
    """segments : liste de tuples (text, dict_style)"""
    for text, style in segments:
        run = p.add_run(text)
        run_style(run, **style)


def add_logo(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(4.5))


def add_banner(doc, badge_text, title_text, subtitle_text):
    # Badge LANCEMENT OFFICIEL
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    paragraph_shade(p, "33CC00")
    run = p.add_run(f"  ▸ {badge_text}  ")
    run_style(run, bold=True, color=COLOR_WHITE, size=10)
    # Titre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    paragraph_shade(p, COLOR_HEADER_BG)
    run = p.add_run(f"  {title_text}  ")
    run_style(run, bold=True, color=COLOR_WHITE, size=18)
    # Sous-titre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    paragraph_shade(p, COLOR_HEADER_BG)
    run = p.add_run(f"  {subtitle_text}  ")
    run_style(run, italic=True, color=RGBColor(0xA5, 0xD6, 0xA7), size=10)


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    paragraph_border_left(p, "33CC00", sz=24)
    run = p.add_run(text.upper())
    run_style(run, bold=True, color=COLOR_GREEN_DARK, size=11)


def add_check_item(doc, text_segments):
    """text_segments : list of (text, is_bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    # Check vert
    run = p.add_run("✓ ")
    run_style(run, bold=True, color=COLOR_GREEN_BRIGHT, size=12)
    # Texte
    for text, is_bold in text_segments:
        run = p.add_run(text)
        run_style(run, bold=is_bold, color=COLOR_TEXT, size=10.5)


def add_sites_box(doc, header_text, sites):
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_before = Pt(4)
    p_hdr.paragraph_format.space_after = Pt(2)
    p_hdr.paragraph_format.left_indent = Cm(0.3)
    paragraph_shade(p_hdr, COLOR_GREEN_LIGHT)
    run = p_hdr.add_run(f"  {header_text}  ")
    run_style(run, bold=True, color=COLOR_GREEN_DARK, size=11)
    for i, site in enumerate(sites, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(2)
        paragraph_shade(p, COLOR_GREEN_LIGHT)
        run = p.add_run(site)
        run_style(run, bold=True, color=COLOR_TEXT_DARK, size=10.5)


def add_why_box(doc, label, text):
    # Label
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    paragraph_shade(p, COLOR_YELLOW_BG)
    paragraph_border_left(p, "F9A825", sz=24)
    run = p.add_run(f"  {label}  ")
    run_style(run, bold=True, color=COLOR_YELLOW_DARK, size=10)
    # Text
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    paragraph_shade(p, COLOR_YELLOW_BG)
    paragraph_border_left(p, "F9A825", sz=24)
    run = p.add_run(f"  {text}  ")
    run_style(run, italic=True, color=COLOR_YELLOW_DARK, size=11)


def add_calendar_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (week, action) in enumerate(rows):
        cell_w = table.rows[i].cells[0]
        cell_a = table.rows[i].cells[1]
        cell_w.width = Cm(3.5)
        cell_a.width = Cm(13)
        cell_w.text = ""
        cell_a.text = ""
        pw = cell_w.paragraphs[0]
        pw.paragraph_format.space_before = Pt(3)
        pw.paragraph_format.space_after = Pt(3)
        run_w = pw.add_run(week)
        run_style(run_w, bold=True, color=COLOR_GREEN_DARK, size=10.5)
        pa = cell_a.paragraphs[0]
        pa.paragraph_format.space_before = Pt(3)
        pa.paragraph_format.space_after = Pt(3)
        run_a = pa.add_run(action)
        run_style(run_a, color=COLOR_TEXT, size=10.5)
        shade(cell_w, COLOR_GREEN_BG)
        shade(cell_a, COLOR_GREEN_BG)
        borders(cell_w, "C8E6C9")
        borders(cell_a, "C8E6C9")


def add_numbered_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    # Numero en rond vert
    run_num = p.add_run(f" {num} ")
    run_style(run_num, bold=True, color=COLOR_WHITE, size=11)
    # Background du num
    rPr = run_num._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "0A4800")
    rPr.append(shd)
    # Texte
    run = p.add_run(f"   {text}")
    run_style(run, color=COLOR_TEXT, size=10.5)


def add_cta_block(doc, text, mailto_url):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    paragraph_shade(p, COLOR_HEADER_BG)
    run = p.add_run(f"  {text}  ")
    run_style(run, color=RGBColor(0xC8, 0xE6, 0xC9), size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    paragraph_shade(p, COLOR_HEADER_BG)
    run = p.add_run("  → Proposer un créneau (egreencitys@gmail.com)  ")
    run_style(run, bold=True, color=COLOR_WHITE, size=11)


def add_footer_block(doc, op):
    # Separateur
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:color"), "33CC00")
    p_bdr.append(top)
    p_pr.append(p_bdr)

    add_para(
        doc, f"EGREENCITY'S — Bornes de recharge électrique en Guyane",
        bold=True, color=COLOR_GREY, size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=2,
    )
    add_para(
        doc, f"{op['adresse']} — {op['code_postal']} {op['ville']} — Guyane française",
        color=COLOR_GREY, size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=2,
    )
    add_para(
        doc, f"Tél : {op['telephone']} | {op['email']} | {op['site_web']}",
        color=COLOR_GREY, size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=2,
    )
    add_para(
        doc, f"SAS au capital de 250 EUR — RCS Cayenne {op['siren']}",
        color=COLOR_GREY, size=8, italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=0,
    )


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT


# ============================================================
#  Generation principale
# ============================================================
def slugify(s):
    return s.lower().replace(" ", "-").replace("é", "e").replace("è", "e") \
            .replace("ê", "e").replace("à", "a").replace("'", "").replace("--", "-")


def generate_docx(mairie, op):
    nb_pdc = mairie["nb_pdc"]
    nb_stations = mairie["nb_stations"]
    s_pdc = "s" if nb_pdc > 1 else ""
    s_station = "s" if nb_stations > 1 else ""

    doc = Document()
    configure_document(doc)
    add_logo(doc)

    # Bandeau d'en-tete
    add_banner(
        doc,
        badge_text="LANCEMENT OFFICIEL",
        title_text=f"Bornes de recharge ADVENIR à {mairie['commune']}",
        subtitle_text="EGREENCITY'S — Signature de convention et calendrier de mise en service",
    )

    # Salutation
    add_para(doc, mairie["civilite_maire"] + ",",
             color=COLOR_TEXT_DARK, size=11, before=8, after=12)

    # Intro
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing = 1.4
    add_runs_mixed(p, [
        ("Faisant suite à nos précédents échanges concernant le projet ADVENIR EGREENCITY'S, "
         "j'ai le plaisir de vous annoncer le ",
         dict(color=COLOR_TEXT, size=10.5)),
        ("lancement officiel",
         dict(bold=True, color=COLOR_GREEN_DARK, size=10.5)),
        (" de notre programme de déploiement de bornes de recharge pour véhicules "
         "électriques en Guyane.",
         dict(color=COLOR_TEXT, size=10.5)),
    ])

    # Etapes sécurisées
    add_section_title(doc, "Étapes sécurisées à ce jour")
    add_check_item(doc, [("Levée de fonds Seed (300 000 €) bouclée", True),
                          (" — closing financier acté", False)])
    add_check_item(doc, [("Subvention ADVENIR confirmée à 75,2 %", True),
                          (" (37 200 € pour les 20 PDC du réseau)", False)])
    add_check_item(doc, [("Devis fournisseur E-TOTEM signé", True),
                          (" (DEV26000037 du 30/04/2026)", False)])
    add_check_item(doc, [("10 bornes commandées", True),
                          (" (20 points de charge) — production lancée", False)])
    add_check_item(doc, [("Délai de livraison Guyane : ", False),
                          ("8 à 10 semaines", True),
                          (" après signature des conventions", False)])
    add_check_item(doc, [("Équipe technique IRVE Qualifelec en cours de recrutement", False)])

    # Ce que nous proposons
    add_section_title(doc, f"Ce que nous vous proposons pour {mairie['commune']}")
    add_sites_box(
        doc,
        f"{nb_pdc} point{s_pdc} de charge sur {nb_stations} station{s_station} :",
        mairie["sites_proposes"],
    )

    add_check_item(doc, [("Modèle ", False),
                          ("E-TOTEM e-Premium AC 2x22 kW", True),
                          (" — fabrication française, peinture milieu humide tropicale", False)])
    add_check_item(doc, [("Logo de la commune sérigraphié sur la borne (visibilité permanente)", False)])
    add_check_item(doc, [("Tarif préférentiel agents municipaux (-50 %) + 2 badges RFID gratuits véhicules de service", False)])
    add_check_item(doc, [("Reporting trimestriel de fréquentation et données énergétiques", False)])
    add_check_item(doc, [("Conformité ", False),
                          ("AFIR 2024", True),
                          (" garantie — sans coût pour la commune", False)])

    # Pourquoi commune ?
    add_why_box(doc, f"Pourquoi {mairie['commune']} ?", mairie["argument_specifique"])

    # Calendrier
    add_section_title(doc, "Calendrier proposé")
    add_calendar_table(doc, [
        ("Semaine 1-2", "Signature de la convention d'occupation du domaine public"),
        ("Semaine 3-4", "Visite technique conjointe (services techniques + EGREENCITY'S)"),
        ("Semaine 5-8", "Travaux de génie civil et raccordement EDF SEI"),
        ("Semaine 9-12", "Livraison, installation et mise en service de la borne"),
        ("Semaine 13", "Inauguration officielle (date à convenir avec vous)"),
    ])

    # Prochaine etape
    add_section_title(doc, "Prochaine étape — RDV de signature")
    add_para(doc,
             "Pour formaliser notre partenariat, je vous propose un rendez-vous de "
             "30 minutes en mairie ou en visioconférence afin de :",
             color=COLOR_TEXT, size=10.5, after=8)
    add_numbered_step(doc, "1", "Signer la convention d'occupation gratuite du domaine public")
    add_numbered_step(doc, "2", "Valider ensemble l'emplacement précis du ou des sites")
    add_numbered_step(doc, "3", "Présenter la fiche technique des bornes à vos services techniques")
    add_numbered_step(doc, "4", "Convenir de la date d'inauguration")

    # CTA
    add_cta_block(
        doc,
        f"Pourriez-vous nous indiquer 2 ou 3 créneaux possibles dans les 2 prochaines "
        f"semaines pour un rendez-vous en mairie ?",
        mailto_url=f"egreencitys@gmail.com",
    )

    # Signature
    add_para(doc,
             f"Vous remerciant par avance de l'attention que vous voudrez bien porter "
             f"à ce projet structurant pour la Guyane et pour {mairie['commune']},",
             color=COLOR_TEXT, size=10.5, before=8, after=8)
    add_para(doc, "Cordialement,", color=COLOR_TEXT, size=10.5, after=12)
    add_para(doc, op["president"], bold=True, color=COLOR_GREEN_DARK, size=12, after=2)
    add_para(doc, "Président — Fondateur, EGREENCITY'S SAS",
             color=COLOR_TEXT, size=10, after=2)
    add_para(doc, f"Tél : {op['telephone']} | {op['email']}",
             color=COLOR_TEXT, size=10, after=0)

    # Footer
    add_footer_block(doc, op)

    return doc


def main():
    with DATA_FILE.open(encoding="utf-8-sig") as f:
        data = json.load(f)
    OP = data["operateur"]

    print(f"=== Generation de {len(data['mairies'])} emails de lancement .docx ===\n")

    for m in data["mairies"]:
        doc = generate_docx(m, OP)
        fname = f"Email_Lancement_Mairie_{m['id']}_{slugify(m['commune'])}.docx"
        out_path = OUT_DIR / fname
        doc.save(out_path)
        size_kb = out_path.stat().st_size / 1024
        print(f"  [OK] {fname}  ({size_kb:.0f} KB)")

    print(f"\n=== Termine : {len(data['mairies'])} fichiers Word generes ===")
    print(f"Dossier : {OUT_DIR}")


if __name__ == "__main__":
    main()
