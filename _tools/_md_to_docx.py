"""
EGREENCITY'S — Conversion Markdown -> Word (.docx)
====================================================

Convertit tous les .md d'un dossier (ou un fichier specifique) en document Word
editable, avec mise en forme charte EGREENCITY'S (charte verte).

USAGE
-----
    python _tools/_md_to_docx.py                                 # dossier ADVENIR
    python _tools/_md_to_docx.py _dossiers/02_Dossier_ADVENIR    # un dossier
    python _tools/_md_to_docx.py path/to/file.md                 # un fichier
    python _tools/_md_to_docx.py --all                           # tous les .md

DEPENDANCES
-----------
    pip install markdown python-docx beautifulsoup4 lxml
"""
from __future__ import annotations
import sys
import re
import argparse
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TARGET = ROOT / "_dossiers" / "02_Dossier_ADVENIR"
LOGO_PATH = ROOT / "logo.png"

# ============================================================
#  Charte EGREENCITY'S
# ============================================================
COLOR_GREEN_DARK   = RGBColor(0x0A, 0x48, 0x00)  # H1, H2, H3
COLOR_GREEN_BRIGHT = RGBColor(0x33, 0xCC, 0x00)  # Accents
COLOR_BLUE         = RGBColor(0x2B, 0x4D, 0xB5)  # Italique / H4
COLOR_TEXT         = RGBColor(0x1A, 0x3A, 0x00)  # Corps
COLOR_GREY         = RGBColor(0x88, 0x88, 0x88)  # Footer
COLOR_LIGHT_BG     = "F4FFF0"  # Fond ligne paire tableau
COLOR_DARK_BG      = "0A4800"  # Fond entete tableau

FONT_NAME = "Calibri"

# Caracteres speciaux a remplacer (memes que pour le PDF)
CHAR_REPLACE = [
    ("☐", "[ ]"), ("☑", "[X]"), ("☒", "[X]"),
    ("✅", "[OK]"), ("❌", "[KO]"), ("✔", "[OK]"), ("✗", "[KO]"), ("✓", "v"),
    ("🟢", "[+]"), ("🟡", "[~]"), ("🔴", "[!]"), ("⚪", "[ ]"), ("⚫", "[#]"),
    ("→", "->"), ("←", "<-"), ("↔", "<->"), ("⇒", "=>"), ("➡", "->"),
    ("≥", ">="), ("≤", "<="), ("≠", "!="), ("≈", "~"),
    ("⚡", "[E]"), ("🔌", "[P]"), ("📍", "[L]"), ("📞", "[T]"), ("📧", "[@]"),
    ("✉", "[@]"), ("📋", ""), ("📁", ""), ("📊", ""), ("📅", ""),
    ("💰", ""), ("💡", ""), ("🔍", ""), ("🎯", ""),
    ("🌐", "Web :"), ("🇫🇷", "FR"), ("🇬🇫", "GF"),
    ("₀", "0"), ("₁", "1"), ("₂", "2"), ("₃", "3"), ("₄", "4"),
    ("₅", "5"), ("₆", "6"), ("₇", "7"), ("₈", "8"), ("₉", "9"),
    ("⁰", "0"), ("¹", "1"), ("²", "2"), ("³", "3"), ("⁴", "4"),
    ("⁵", "5"), ("⁶", "6"), ("⁷", "7"), ("⁸", "8"), ("⁹", "9"),
    ("ᵉ", "e"), ("ʳ", "r"), ("ⁿ", "n"), ("ᵃ", "a"), ("ᵒ", "o"),
    (" ", " "), (" ", " "),
]


def shade_cell(cell, color_hex: str):
    """Applique une couleur de fond a une cellule de tableau."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """Ajoute des bordures fines grises a une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_horizontal_rule(doc):
    """Insere une ligne horizontale verte (HR markdown)."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "33CC00")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_run_style(run, *, bold=False, italic=False, color=COLOR_TEXT,
                  size=11, font_name=FONT_NAME, underline=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if underline:
        run.font.underline = True
    # Compatibilite ASCII / EastAsian
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_runs_from_inline(paragraph, element, *, base_bold=False,
                         base_italic=False, base_color=COLOR_TEXT,
                         base_size=11):
    """Parcourt recursivement les enfants d'un element pour ajouter des runs
    avec gestion bold/italic/code/links."""
    if isinstance(element, NavigableString):
        text = str(element)
        if not text:
            return
        run = paragraph.add_run(text)
        set_run_style(run, bold=base_bold, italic=base_italic,
                      color=base_color, size=base_size)
        return

    name = element.name
    if name in ("strong", "b"):
        for child in element.children:
            add_runs_from_inline(paragraph, child, base_bold=True,
                                 base_italic=base_italic,
                                 base_color=COLOR_GREEN_DARK,
                                 base_size=base_size)
    elif name in ("em", "i"):
        for child in element.children:
            add_runs_from_inline(paragraph, child, base_bold=base_bold,
                                 base_italic=True, base_color=COLOR_BLUE,
                                 base_size=base_size)
    elif name == "code":
        run = paragraph.add_run(element.get_text())
        set_run_style(run, bold=base_bold, italic=base_italic,
                      color=COLOR_GREEN_DARK, size=base_size,
                      font_name="Consolas")
        # Fond gris clair via shading
        rPr = run._element.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F0F4EC")
        rPr.append(shd)
    elif name == "a":
        text = element.get_text()
        href = element.get("href", "")
        run = paragraph.add_run(text)
        set_run_style(run, bold=base_bold, italic=base_italic,
                      color=COLOR_GREEN_DARK, size=base_size,
                      underline=True)
        # Note : on n'ajoute pas un vrai hyperlien Word pour rester simple ;
        # le texte conserve l'URL si elle est differente
        if href and href != text:
            extra = paragraph.add_run(f" ({href})")
            set_run_style(extra, italic=True, color=COLOR_GREY,
                          size=base_size - 1)
    elif name == "br":
        paragraph.add_run().add_break()
    elif name in ("span",):
        # Champ a remplir : cas .fillin ou .fillin-short
        cls = element.get("class", [])
        if cls and ("fillin" in cls or "fillin-short" in cls):
            # Insere une serie d'underscores stylises (champ a remplir)
            length = 30 if "fillin" in cls and "fillin-short" not in cls else 12
            run = paragraph.add_run("_" * length)
            set_run_style(run, color=COLOR_GREY, size=base_size)
        else:
            for child in element.children:
                add_runs_from_inline(paragraph, child, base_bold=base_bold,
                                     base_italic=base_italic,
                                     base_color=base_color,
                                     base_size=base_size)
    else:
        for child in element.children:
            add_runs_from_inline(paragraph, child, base_bold=base_bold,
                                 base_italic=base_italic,
                                 base_color=base_color,
                                 base_size=base_size)


def add_heading(doc, text, level):
    """Ajoute un titre stylise EGREENCITY'S."""
    if level == 1:
        # H1 : fond vert fonce, texte blanc
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(14)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "0A4800")
        p_pr.append(shd)
        # Bordure gauche verte
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), "33CC00")
        p_bdr.append(left)
        p_pr.append(p_bdr)
        # Padding interne via indentation
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(text)
        set_run_style(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=18)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        # Bordure bas verte
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "33CC00")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        run = p.add_run(text)
        set_run_style(run, bold=True, color=COLOR_GREEN_DARK, size=14)
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_style(run, bold=True, color=COLOR_GREEN_DARK, size=12)
    elif level == 4:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        set_run_style(run, bold=True, color=COLOR_BLUE, size=11)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_style(run, bold=True, color=COLOR_GREEN_DARK, size=10)


def add_paragraph_with_inlines(doc, element):
    """Ajoute un paragraphe avec gestion des runs (bold/italic/code/links)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    for child in element.children:
        add_runs_from_inline(p, child)
    return p


def add_list(doc, list_element, ordered=False, level=0):
    """Ajoute une liste ordonnee ou non. Les sous-listes sont gerees."""
    style = "List Number" if ordered else "List Bullet"
    for li in list_element.find_all("li", recursive=False):
        # Texte direct du li (sans les sous-listes)
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        if level > 0:
            p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
        # Filtrer les enfants qui ne sont pas des listes
        for child in li.children:
            if isinstance(child, NavigableString):
                add_runs_from_inline(p, child)
            elif child.name in ("ul", "ol"):
                continue  # Traite plus bas
            else:
                add_runs_from_inline(p, child)
        # Sous-listes
        for sublist in li.find_all(["ul", "ol"], recursive=False):
            add_list(doc, sublist, ordered=(sublist.name == "ol"),
                     level=level + 1)


def add_blockquote(doc, element):
    """Ajoute une citation stylisee (bordure gauche bleue, italique)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    # Bordure gauche bleue
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2B4DB5")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    # Fond bleu clair
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F3FB")
    p_pr.append(shd)
    # Contenu en italique
    text = element.get_text(" ", strip=True)
    run = p.add_run(text)
    set_run_style(run, italic=True, color=COLOR_TEXT, size=10)


def add_code_block(doc, element):
    """Ajoute un bloc de code (fond clair, monospace)."""
    text = element.get_text()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Bordure gauche verte vive + fond clair
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "33CC00")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4FFF0")
    p_pr.append(shd)
    run = p.add_run(text.rstrip())
    set_run_style(run, color=COLOR_TEXT, size=9.5, font_name="Consolas")


def add_table(doc, table_element):
    """Ajoute un tableau Word avec entete vert fonce + lignes alternees."""
    # Compter les lignes
    rows = table_element.find_all("tr")
    if not rows:
        return
    # Premiere ligne = entete (THEAD ou premiere TR)
    n_cols = max(len(r.find_all(["td", "th"])) for r in rows)
    if n_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        is_header = (r_idx == 0 and any(c.name == "th" for c in cells)) or \
                    (tr.parent.name == "thead")

        for c_idx, cell_el in enumerate(cells):
            if c_idx >= n_cols:
                break
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Reset paragraph existant
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Determiner alignement (markdown : | --- | :--- | ---: | :---: |)
            # Par defaut a gauche pour text, droite pour nombres
            for child in cell_el.children:
                add_runs_from_inline(
                    p, child,
                    base_bold=is_header,
                    base_color=(RGBColor(0xFF, 0xFF, 0xFF) if is_header else COLOR_TEXT),
                    base_size=10 if is_header else 9.5,
                )
            # Style cellule
            set_cell_borders(cell)
            if is_header:
                shade_cell(cell, COLOR_DARK_BG)
            elif r_idx % 2 == 1:
                shade_cell(cell, COLOR_LIGHT_BG)


def add_header_section(doc, title):
    """En-tete du document : logo EGREENCITY'S + tagline.
    Logo : width fixe (preserve le ratio natif, pas d'etirement)."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        run = p.add_run()
        try:
            # Fixe la LARGEUR (pas la hauteur) pour preserver le ratio natif
            # Logo natif ~600x150 px (ratio 4:1) -> width 4.5cm = height ~1.1cm
            run.add_picture(str(LOGO_PATH), width=Cm(4.5))
        except Exception:
            pass
    else:
        run = p.add_run("EGREENCITY'S")
        set_run_style(run, bold=True, color=COLOR_GREEN_BRIGHT, size=14)
        p.add_run("\t\t")
        run2 = p.add_run("Bornes de recharge electrique - Guyane francaise")
        set_run_style(run2, italic=True, color=COLOR_BLUE, size=8)
    # Footer : RCS + page numero
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        "EGREENCITY'S SAS · RCS Cayenne 878 682 854 · "
        "egreencitys@gmail.com · +33 6 51 14 11 18"
    )
    set_run_style(run, color=COLOR_GREY, size=8)


def configure_document(doc):
    """Marges + style par defaut."""
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    # Style Normal
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT


# ============================================================
#  Conversion principale
# ============================================================
def preprocess_md(text: str) -> str:
    for src, dst in CHAR_REPLACE:
        text = text.replace(src, dst)
    return text


def md_to_docx(md_path: Path, out_path: Path | None = None) -> Path:
    md_text = md_path.read_text(encoding="utf-8")
    md_text = preprocess_md(md_text)

    # markdown -> HTML
    html = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists", "attr_list"],
        output_format="html5",
    )
    soup = BeautifulSoup(f"<div>{html}</div>", "html.parser")
    body = soup.find("div")

    # Document
    doc = Document()
    configure_document(doc)
    # Titre = premier H1 du .md
    title = md_path.stem
    for line in md_text.splitlines():
        if line.startswith("# "):
            title = line[2:].strip()
            break
    add_header_section(doc, title)

    # Parcours des elements top-level
    for el in body.children:
        if isinstance(el, NavigableString):
            txt = str(el).strip()
            if txt:
                p = doc.add_paragraph()
                run = p.add_run(txt)
                set_run_style(run, color=COLOR_TEXT, size=10.5)
            continue
        name = el.name
        if name == "h1":
            add_heading(doc, el.get_text(strip=True), 1)
        elif name == "h2":
            add_heading(doc, el.get_text(strip=True), 2)
        elif name == "h3":
            add_heading(doc, el.get_text(strip=True), 3)
        elif name == "h4":
            add_heading(doc, el.get_text(strip=True), 4)
        elif name in ("h5", "h6"):
            add_heading(doc, el.get_text(strip=True), 5)
        elif name == "p":
            add_paragraph_with_inlines(doc, el)
        elif name == "ul":
            add_list(doc, el, ordered=False)
        elif name == "ol":
            add_list(doc, el, ordered=True)
        elif name == "blockquote":
            add_blockquote(doc, el)
        elif name == "pre":
            code = el.find("code")
            add_code_block(doc, code if code else el)
        elif name == "table":
            add_table(doc, el)
        elif name == "hr":
            add_horizontal_rule(doc)
        elif name == "dl":
            # dl.kvlist (tables 2-cols converties)
            for kvrow in el.find_all("div", class_="kvrow"):
                dt = kvrow.find("dt")
                dd = kvrow.find("dd")
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                if dt:
                    run = p.add_run(dt.get_text(strip=True) + "\t")
                    set_run_style(run, bold=True, color=COLOR_GREEN_DARK,
                                  size=10)
                if dd:
                    for child in dd.children:
                        add_runs_from_inline(p, child)
        else:
            # Element inconnu : on traite ses paragraphes enfants
            for child in el.find_all(["p", "ul", "ol", "blockquote",
                                      "table", "hr"], recursive=False):
                if child.name == "p":
                    add_paragraph_with_inlines(doc, child)
                elif child.name in ("ul", "ol"):
                    add_list(doc, child, ordered=(child.name == "ol"))
                elif child.name == "blockquote":
                    add_blockquote(doc, child)
                elif child.name == "table":
                    add_table(doc, child)
                elif child.name == "hr":
                    add_horizontal_rule(doc)

    # Sauvegarde
    if out_path is None:
        out_path = md_path.with_suffix(".docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def collect_md_files(target: Path) -> list[Path]:
    if target.is_file() and target.suffix.lower() == ".md":
        return [target]
    if target.is_dir():
        return sorted(target.rglob("*.md"))
    return []


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Convertit des .md en .docx (charte EGREENCITY'S)."
    )
    ap.add_argument("target", nargs="?", default=str(DEFAULT_TARGET),
                    help="Fichier .md ou dossier (defaut : 02_Dossier_ADVENIR)")
    ap.add_argument("--all", action="store_true",
                    help="Convertir tous les .md du projet")
    ap.add_argument("-o", "--out", help="Repertoire de sortie")
    args = ap.parse_args()

    if args.all:
        md_files = sorted(ROOT.rglob("*.md"))
        skip = {".git", "_archive", "__pycache__", "node_modules"}
        md_files = [p for p in md_files if not any(s in p.parts for s in skip)]
    else:
        target = Path(args.target)
        if not target.is_absolute():
            target = ROOT / target
        md_files = collect_md_files(target)

    if not md_files:
        print(f"[ERREUR] Aucun .md trouve.")
        return 1

    out_dir = Path(args.out) if args.out else None
    if out_dir and not out_dir.is_absolute():
        out_dir = ROOT / out_dir

    print(f"=== Conversion {len(md_files)} fichier(s) en .docx ===\n")
    ok = errors = 0
    for md in md_files:
        try:
            target_docx = (out_dir / (md.stem + ".docx")) if out_dir \
                else md.with_suffix(".docx")
            md_to_docx(md, target_docx)
            size_kb = target_docx.stat().st_size / 1024
            rel = target_docx.relative_to(ROOT)
            print(f"  [OK]  {rel}  ({size_kb:.0f} KB)")
            ok += 1
        except Exception as e:
            print(f"  [ERR] {md.name}: {e}")
            import traceback
            traceback.print_exc()
            errors += 1

    print(f"\n=== Termine : {ok} OK, {errors} erreur(s) ===")
    return 0 if errors == 0 else 2


if __name__ == "__main__":
    sys.exit(main())
